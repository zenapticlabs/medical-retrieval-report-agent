import json
import os
import requests
import msal
import logging
import boto3
import tempfile
import zipfile
from datetime import datetime
from typing import Dict, List, Any
from io import BytesIO
import PyPDF2
from docx import Document
import re
import concurrent.futures
import tiktoken
import time
import google.generativeai as genai

# Set up logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Environment variables
TENANT_ID = os.environ.get('TENANT_ID')
CLIENT_ID = os.environ.get('CLIENT_ID')
CLIENT_SECRET = os.environ.get('CLIENT_SECRET')
SITE_ID = os.environ.get('SITE_ID')
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
GEMINI_API_KEY = "AIzaSyDmYaWkrqGaaH6ZEldo3OLQFNuD4UAmQLg"
S3_BUCKET = os.environ.get('S3_BUCKET', 'medical-chronology')

# Initialize S3 client
s3_client = boto3.client('s3')

# Master prompt for medical chronology
MASTER_PROMPT = """System (role = "expert medical-records analyst")
You are a licensed U.S. nurse consultant who prepares comprehensive medical chronology summaries for litigation support.
Your task is to extract, normalize and insert information from ALL files provided in this chat (PDF, DOCX, etc.) into a detailed medical chronology format.

Governing rules:
- Confidentiality – treat all PHI as privileged. Do not add commentary, opinions or legal conclusions.
- Date format – MM/DD/YYYY. If only month-year is available, use "MM/YYYY".
- Time ordering – all chronologies must be in strict ascending date order.
- Consistency – spell provider & facility names exactly as they first appear; thereafter use that spelling.
- Bates numbers – Use the actual file name with properly formatted batch numbers and page references. 
  * For files with batch numbers like "UIH_178", "UCM_321", "PMH_156", format as: "FileName_BatchNumber_PageNumber"
  * Examples: "Adkisson_Patricia_2019.10.31_PMH_UIH_178_p1", "PathologyReport_UCM_321_p2"
  * For date-based batches like "2019.10.31_178", format as: "FileName_DATE_2019.10.31_178_p1"
  * For ranges like "UIH_178-185", use the specific batch number for the page: "FileName_UIH_178_p1", "FileName_UIH_179_p1", etc.
  * ALWAYS include the page number in every reference (e.g., ..._p1, ..._p2, etc.), even for single-page files
  * For the Word document output, use the format: "DOCUMENT_NAME-Page X" (e.g., "Adkisson Patricia 2019.10.31 PMH UIH 178-Page 1")
- Abbreviations – expand the first mention (e.g., "CT (computed tomography)") then abbreviate thereafter.
- Objectivity – copy pertinent findings verbatim or paraphrase neutrally; avoid diagnostic speculation.

Required sections to extract and organize:

1. META TABLE:
   - Patient Name
   - Date of Birth
   - Date of Diagnosis
   - Age at Diagnosis
   - Primary Diagnosis/Track 1 Disease
   - Attorney (if mentioned)
   - Preparation Date

2. SUMMARY OF MEDICAL HISTORY:
   - Date of Birth
   - Date of Diagnosis  
   - Age at Diagnosis
   - T1 Injury + Other Related Secondary Diagnoses
   - Chemotherapy Treatments (list all with dates)
   - Detailed chronological narrative of key events (diagnosis, treatments, follow-ups)
   - Sufferings (patient-reported symptoms and impacts)
   - Other medical history (prior to T1 and after T1)
   - Risk Factors
   - Additional Information (Family Hx, Social Hx, Work history)
   - Missing records (if any)

3. CHRONOLOGICAL MEDICAL RECORD REVIEW:
   - One row per encounter/result with 4 columns:
     Date | Medical Provider/Medical Facility | Summary | Bates #
   - Summaries should be DETAILED (50-150 words each) including:
     * Chief complaint or reason for visit
     * Key findings from physical exam
     * Diagnostic test results with specific values
     * Treatment plans and medications prescribed
     * Follow-up recommendations
     * Any significant clinical observations
   - Use present tense ("Reports nausea…", "CT shows…")
   - Include specific lab values, vital signs, and measurements when available
   - Extract ALL relevant medical information from each document
   - For Bates numbers, use properly formatted batch information with actual file names and page references

4. RECORD INDEX:
   - One row per distinct facility/medical provider
   - Show continuous Bates range using properly formatted batch numbers (e.g., "Adkisson_Patricia_UIH_178-185", "PathologyReport_UCM_321-325")
   - Include overall date span for each facility
   - Include detailed description of record types (e.g., "Oncology visits, pathology reports, radiology studies, lab results")
   - Group by facility and provide comprehensive coverage of all reviewed records

5. QUALIFYING DIAGNOSIS TABLE:
  - For each qualifying cancer or catastrophic diagnosis, create a separate row with these columns:
    * Diagnosis: (concise summary, e.g., “High-grade serous carcinoma, tubo-ovarian, FIGO IIIB” — do NOT include full pathology text or all details, just the main diagnosis and stage)
    * Dx Reference: (file and page or Bates reference where diagnosis appears)
    * Treatment: (treatments given, with dates, in concise bullet or list format)
    * Tx Reference: (file and page or Bates reference where treatment evidence appears)
  - Only include diagnoses that are serious, life-threatening, or may qualify for litigation/insurance/benefits (e.g., any cancer diagnosis with pathologic/clinical proof, or catastrophic injury).
  - Be precise and specific. Each reference must point directly to the page or Bates number where the diagnosis or treatment is proven.
  - If multiple sources exist for Dx/Tx, list all significant references.
  - Format as a JSON array under "qualifying_diagnosis_table".
  - The “diagnosis” field should be brief and clinically focused (1-2 lines), not a full pathology report or list of all findings.

IMPORTANT: Be extremely thorough in extracting information. Each chronological entry should be comprehensive and detailed, not just 2-3 words. Include all relevant medical findings, test results, treatments, and clinical observations from the source documents. Always reference the actual file names with properly formatted batch numbers in Bates references.

Please analyze all the provided medical files and return a structured JSON response with the following format:

{
  "meta_table": {
    "patient_name": "",
    "date_of_birth": "",
    "date_of_diagnosis": "",
    "age_at_diagnosis": "",
    "primary_diagnosis": "",
    "attorney": "",
    "preparation_date": ""
  },
  "medical_history_summary": {
    "date_of_birth": "",
    "date_of_diagnosis": "",
    "age_at_diagnosis": "",
    "t1_injury_and_secondary": "",
    "chemotherapy_treatments": [],
    "chronological_narrative": "",
    "sufferings": [],
    "other_medical_history": {
      "prior_to_t1": [],
      "after_t1": []
    },
    "risk_factors": [],
    "additional_information": {
      "family_history": "",
      "social_history": "",
      "work_history": ""
    },
    "missing_records": ""
  },
  "chronological_records": [
    {"date": "MM/DD/YYYY", "provider_facility": "", "summary": "DETAILED SUMMARY WITH SPECIFIC FINDINGS, LAB VALUES, TREATMENTS, AND CLINICAL OBSERVATIONS (50-150 words)", "bates": "ActualFileName_ProperlyFormattedBatchNumber_PageNumber"}
  ],
  "record_index": [
    {"facility": "", "bates_range": "ActualFileName_ProperlyFormattedBatchNumber_StartPage-EndPage", "date_range": "", "description": "DETAILED DESCRIPTION OF RECORD TYPES AND CONTENT"}
  ],
  "qualifying_diagnosis_table": [
    {
      "diagnosis": "",
      "dx_reference": "",
      "treatment": "",
      "tx_reference": ""
    }
  ],
}

Begin comprehensive analysis of the provided medical files now. Extract ALL available medical information and create detailed, thorough entries using actual file names with properly formatted batch numbers for Bates references."""

def get_access_token() -> str:
    """
    Get access token using MSAL for SharePoint access
    """
    try:
        app = msal.ConfidentialClientApplication(
            client_id=CLIENT_ID,
            client_credential=CLIENT_SECRET,
            authority=f"https://login.microsoftonline.com/{TENANT_ID}"
        )
        
        result = app.acquire_token_for_client(
            scopes=["https://graph.microsoft.com/.default"]
        )
        
        token = result.get("access_token")
        if not token:
            logger.error(f"Token acquisition failed: {result}")
            raise Exception(f"Token error: {result}")
        
        return token
    except Exception as e:
        logger.error(f"Error getting access token: {str(e)}")
        raise

def list_children(path: str, token: str) -> List[Dict]:
    """
    List all items (files and folders) directly under the given path
    """
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json"
    }
    
    # Construct the URL for the SharePoint path
    if path:
        url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drive/root:/{path}:/children"
    else:
        url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drive/root/children"
    
    all_items = []
    
    while url:
        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            data = response.json()
            
            for item in data.get("value", []):
                all_items.append(item)
            
            # Handle pagination
            url = data.get("@odata.nextLink")
            
        except requests.exceptions.RequestException as e:
            logger.error(f"Error fetching children for path '{path}': {str(e)}")
            raise
    
    return all_items

def walk_sharepoint_path(path: str, token: str, max_depth: int = 5, current_depth: int = 0) -> List[Dict]:
    """
    Recursively walk through SharePoint folders and collect all files and folders
    """
    if current_depth >= max_depth:
        logger.warning(f"Maximum depth ({max_depth}) reached for path: {path}")
        return []
    
    result = []
    
    try:
        items = list_children(path, token)
        
        for item in items:
            item_name = item["name"]
            
            # Skip files with "_CHR Claim" substring
            if "_CHR Claim" in item_name:
                logger.info(f"Skipping file with '_CHR Claim': {item_name}")
                continue
            
            item_info = {
                "name": item_name,
                "path": f"{path}/{item_name}" if path else item_name,
                "download_url": item.get("@microsoft.graph.downloadUrl"),
                "size": item.get("size", 0),
                "type": "folder" if "folder" in item else "file",
                "depth": current_depth,
                "lastModified": item.get("lastModifiedDateTime"),
                "created": item.get("createdDateTime")
            }
            
            if "folder" in item:
                # It's a folder
                item_info["childCount"] = item["folder"].get("childCount", 0)
                result.append(item_info)
                
                # Recursively get children if it's a folder
                sub_path = f"{path}/{item_name}" if path else item_name
                child_items = walk_sharepoint_path(sub_path, token, max_depth, current_depth + 1)
                result.extend(child_items)
            else:
                # It's a file
                result.append(item_info)
                
    except Exception as e:
        logger.error(f"Error walking path '{path}': {str(e)}")
        # Don't re-raise, just log and continue
    
    return result

def download_file_content(download_url: str, file_name: str) -> list:
    """Download and extract text content from SharePoint file, paginated by page_number."""
    try:
        response = requests.get(download_url)
        response.raise_for_status()
        
        file_extension = os.path.splitext(file_name)[1].lower()
        content = []
        
        if file_extension == '.pdf':
            # Extract text from PDF page by page
            pdf_reader = PyPDF2.PdfReader(BytesIO(response.content))
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text() or ""
                content.append({"page_number": i + 1, "text": page_text})
        
        elif file_extension in ['.docx', '.doc']:
            # Extract text from Word document (treat as single page)
            doc = Document(BytesIO(response.content))
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            content.append({"page_number": 1, "text": full_text})
        
        elif file_extension == '.txt':
            # Plain text file (treat as single page)
            text = response.content.decode('utf-8', errors='ignore')
            content.append({"page_number": 1, "text": text})
        
        else:
            logger.warning(f"Unsupported file type: {file_extension} for file: {file_name}")
            content.append({"page_number": 1, "text": f"[Unsupported file type: {file_extension}]"})
        
        return content
    
    except Exception as e:
        logger.error(f"Error downloading/processing file {file_name}: {str(e)}")
        return [{"page_number": 1, "text": f"[Error processing file: {str(e)}]"}]

def extract_batch_info(file_name: str) -> str:
    """
    Extract and format batch information from file names.
    Handles various batch number formats commonly found in medical records.
    """
    # Convert to uppercase for consistent matching
    file_upper = file_name.upper()
    
    # Pattern matching for different batch formats
    batch_patterns = [
        # UIH batch numbers (e.g., UIH_178, UIH_178-185, UIH_178_p1)
        (r'UIH[_\s]*(\d+)(?:[-_](\d+))?(?:_P(\d+))?', 'UIH'),
        # UCM batch numbers (e.g., UCM_321, UCM_321-325, UCM_321_p2)
        (r'UCM[_\s]*(\d+)(?:[-_](\d+))?(?:_P(\d+))?', 'UCM'),
        # PMH batch numbers (e.g., PMH_156, PMH_156-160)
        (r'PMH[_\s]*(\d+)(?:[-_](\d+))?(?:_P(\d+))?', 'PMH'),
        # Generic batch numbers (e.g., BATCH_001, B001-B005)
        (r'(?:BATCH[_\s]*|B)(\d+)(?:[-_](?:BATCH[_\s]*|B)?(\d+))?(?:_P(\d+))?', 'BATCH'),
        # Medical records with numbers (e.g., MR_2019_001, MEDICAL_001-005)
        (r'(?:MR|MEDICAL)[_\s]*(?:\d{4}[_\s]*)?(\d+)(?:[-_](\d+))?(?:_P(\d+))?', 'MR'),
        # Date-based batches (e.g., 2019.10.31_178, 20191031_178-185)
        (r'(\d{4})[.\-_]?(\d{2})[.\-_]?(\d{2})[_\s]*(\d+)(?:[-_](\d+))?(?:_P(\d+))?', 'DATE'),
        # Simple numeric patterns at end of filename (e.g., Document_178, Report_178-185)
        (r'[_\s](\d+)(?:[-_](\d+))?(?:_P(\d+))?(?:\.[A-Z]{2,4})?$', 'DOC')
    ]
    
    for pattern, batch_type in batch_patterns:
        match = re.search(pattern, file_upper)
        if match:
            groups = match.groups()
            
            if batch_type == 'DATE':
                # Handle date-based format
                year, month, day, start_num = groups[:4]
                end_num = groups[4] if len(groups) > 4 and groups[4] else None
                page_num = groups[5] if len(groups) > 5 and groups[5] else None
                
                batch_range = f"{start_num}"
                if end_num:
                    batch_range += f"-{end_num}"
                if page_num:
                    batch_range += f"_p{page_num}"
                
                return f"{batch_type} {year}.{month}.{day}_{batch_range}"
            
            else:
                # Handle standard batch formats
                start_num = groups[0]
                end_num = groups[1] if len(groups) > 1 and groups[1] else None
                page_num = groups[2] if len(groups) > 2 and groups[2] else None
                
                batch_range = start_num
                if end_num:
                    batch_range += f"-{end_num}"
                if page_num:
                    batch_range += f"_p{page_num}"
                
                return f"{batch_type} {batch_range}"
    
    # If no specific pattern found, try to identify facility/source
    if any(facility in file_upper for facility in ['UIH', 'UNIVERSITY', 'HOSPITAL']):
        return "UIH batch (unspecified)"
    elif any(facility in file_upper for facility in ['UCM', 'MEDICAL CENTER']):
        return "UCM batch (unspecified)"
    elif any(facility in file_upper for facility in ['PMH', 'PRESBYTERIAN']):
        return "PMH batch (unspecified)"
    elif any(word in file_upper for word in ['PATHOLOGY', 'PATH']):
        return "Pathology batch"
    elif any(word in file_upper for word in ['RADIOLOGY', 'RAD', 'XRAY', 'CT', 'MRI']):
        return "Radiology batch"
    elif any(word in file_upper for word in ['LAB', 'LABORATORY']):
        return "Lab batch"
    else:
        return "General medical batch"

def format_document_reference(file_name: str, page_number: int) -> str:
    """
    Format document reference for the Word document in the format: DOCUMENT_NAME-Page X
    """
    # Remove file extension
    base_name = os.path.splitext(file_name)[0]
    # Replace underscores and other separators with spaces for readability
    formatted_name = base_name.replace('_', ' ').replace('-', ' ')
    
    # If page_number is 0, return just the document name (for ranges)
    if page_number == 0:
        return formatted_name
    else:
        return f"{formatted_name}-Page {page_number}"

def num_tokens_from_string(string: str, model: str = "gpt-4o") -> int:
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(string))

def batch_files(files_data, max_tokens_per_batch=120000, model="gpt-4o"):
    batches = []
    current_batch = []
    current_tokens = 0

    for file_data in files_data:
        file_content = file_data['content']
        file_name = file_data['name']
        page_number = file_data.get('page_number', 1)
        batch_info = extract_batch_info(file_name)
        file_str = f"\n\n--- FILE: {file_name} ({batch_info}) PAGE {page_number} ---\nContent:\n{file_content}"
        file_tokens = num_tokens_from_string(file_str, model)

        if current_tokens + file_tokens > max_tokens_per_batch and current_batch:
            batches.append(current_batch)
            current_batch = []
            current_tokens = 0

        current_batch.append(file_data)
        current_tokens += file_tokens

    if current_batch:
        batches.append(current_batch)

    return batches

def process_batch(batch, batch_index):
    files_content = ""
    for i, file_data in enumerate(batch, 1):
        file_name = file_data['name']
        page_number = file_data.get('page_number', 1)
        batch_info = extract_batch_info(file_name)
        files_content += f"\n\n--- FILE {i}: {file_name} ({batch_info}) PAGE {page_number} ---\n"
        files_content += f"File Path: {file_data.get('path', 'N/A')}\n"
        files_content += f"Content:\n{file_data['content']}"

    import openai
    openai.api_key = OPENAI_API_KEY
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": MASTER_PROMPT},
            {"role": "user", "content": f"Please analyze these medical files and create a chronology summary. Use the actual file names provided for Bates number references:\n{files_content}"}
        ],
        max_tokens=4000,
        temperature=0.1
    )
    response_content = response['choices'][0]['message']['content']
    return response_content

def process_files_with_gemini(files_data: list, use_mock: bool = False) -> dict:
    """Process file contents with Google Gemini 1.5 Pro to create medical chronology"""
    try:
        if use_mock:
            # Mock response for testing
            return {
                "meta_table": {
                    "patient_name": "Adkisson, Patricia",
                    "date_of_birth": "12/16/1959",
                    "date_of_diagnosis": "11/17/2021",
                    "age_at_diagnosis": "61",
                    "primary_diagnosis": "Bladder Carcinoma",
                    "attorney": "Rosen Injury Lawyers",
                    "preparation_date": datetime.now().strftime("%m/%d/%Y")
                },
                "medical_history_summary": {
                    "date_of_birth": "12/16/1959",
                    "date_of_diagnosis": "11/17/2021",
                    "age_at_diagnosis": "61",
                    "t1_injury_and_secondary": "Bladder carcinoma - Morphologically low grade, mature B-Cell lymphoma, stage IV marginal zone",
                    "chemotherapy_treatments": ["02/10/2022-06/02/2022 R-CVP x 6"],
                    "chronological_narrative": "10/20/2021 CT Thorax revealed right upper lobe medial airspace disease consistent with atelectasis and/or infiltrate associated with multiple bilateral spiculated nodules and masses. 11/17/2021 Surgical Pathology showed Morphologically low-grade, mature B-cell lymphoma, favor marginal zone. 12/16/2021 Diagnosed Acute Lymphoma involving the lung. 01/20/2022 Initial Oncology Consultation with Dr. Nagesh Jayaram for Stage IV Marginal Zone Lymphoma. 02/10/2022-06/02/2022 Completed 6 cycles of R-CVP chemotherapy.",
                    "sufferings": [
                        "Experienced cough with productive sputum for 4-5 months before diagnosis",
                        "Fatigue and worsening peripheral neuropathy",
                        "Developed Monoclonal gammopathy"
                    ],
                    "other_medical_history": {
                        "prior_to_t1": ["HTN", "Sleep Apnea", "Lumbar radiculopathy", "Cardiac murmur"],
                        "after_t1": ["Small fiber sensory poly neuropathy", "Monoclonal gammopathy", "EMG with moderately severe sensorimotor mixed axonal and demyelinating polyneuropathy"]
                    },
                    "risk_factors": ["One sister – died of cancer"],
                    "additional_information": {
                        "family_history": "Both parents deceased. Father – alcohol related issues. Mother – died of complications from hip surgery. One sister – died of cancer.",
                        "social_history": "Never smoker. Has no history of drinking.",
                        "work_history": "Marine corp x 20 years; retired combat tours x2 18mo total Persian gulf and Beirut (06/26/1978-06/30/1998). Police officer x 29 years in Jacksonville."
                    },
                    "missing_records": "None identified"
                },
                "chronological_records": [
                    {"date": "10/31/2019", "provider_facility": "University Iowa Hospital", "summary": "Past medical history documented. Patient presents for routine follow-up. Vital signs: BP 140/90, HR 72, Temp 98.6°F, Weight 185 lbs. Physical examination reveals no acute distress. Cardiovascular exam shows regular rate and rhythm with no murmurs, gallops, or rubs. Respiratory exam shows clear lung fields bilaterally. Abdomen is soft, non-tender, non-distended with no hepatosplenomegaly. Extremities show no edema. Neurological exam is grossly intact. Assessment: Stable chronic conditions. Plan: Continue current medications and follow up in 6 months.", "bates": "Adkisson_Patricia_2019.10.31_PMH_UIH_178_p1"},
                    {"date": "04/01/2021", "provider_facility": "University Medical Center", "summary": "Pathology report for tissue biopsy. Specimen received: Right lung mass biopsy. Gross description: Multiple tan-white tissue fragments measuring 0.8 x 0.6 x 0.3 cm in aggregate. Microscopic examination reveals atypical lymphoid infiltrate with small to medium-sized lymphocytes showing irregular nuclear contours. Immunohistochemical stains show CD20 positive, CD3 negative, CD5 negative, CD10 negative, BCL-2 positive, BCL-6 negative. Ki-67 proliferation index is approximately 15%. Final diagnosis: Atypical lymphoid infiltrate, suspicious for low-grade B-cell lymphoma. Recommend additional studies for definitive classification.", "bates": "Adkisson_Patricia_2021.04.01_PathReport_UCM_321_p1"},
                    {"date": "05/24/2021", "provider_facility": "University Medical Center", "summary": "Follow-up pathology report with additional immunohistochemical studies. Additional stains performed: Cyclin D1 negative, SOX11 negative, CD23 negative, CD43 negative. Flow cytometry analysis shows monoclonal B-cell population with kappa light chain restriction. Molecular studies reveal no evidence of BCL-1 or BCL-2 gene rearrangements. Final diagnosis: Marginal zone lymphoma, low-grade. The neoplastic cells show characteristic immunophenotype and morphology consistent with extranodal marginal zone lymphoma of mucosa-associated lymphoid tissue (MALT) type.", "bates": "Adkisson_Patricia_2021.05.24_PathReport_UCM_365_p2"},
                    {"date": "01/17/2022", "provider_facility": "Medical Center", "summary": "Follow-up oncology visit. Patient reports feeling well with no new symptoms. Physical examination shows no evidence of disease progression. Vital signs stable: BP 138/88, HR 70, Temp 98.4°F. No palpable lymphadenopathy. Chest examination reveals clear lung fields with no wheezes, rales, or rhonchi. Cardiovascular examination shows regular rate and rhythm. Abdomen is soft and non-tender. Extremities show no edema. Laboratory studies show normal CBC with WBC 7.2 K/μL, Hgb 14.2 g/dL, Platelets 245 K/μL. Comprehensive metabolic panel within normal limits. Assessment: No evidence of disease (NED). Plan: Continue surveillance with repeat imaging in 3 months.", "bates": "Adkisson_Patricia_2022.01.17_NED_UCM_146_p1"},
                    {"date": "02/16/2022", "provider_facility": "Medical Center", "summary": "Treatment planning visit in oncology history. Patient presents for discussion of treatment options for Stage IV marginal zone lymphoma. Recent PET/CT scan shows stable disease with no new areas of involvement. Bone marrow biopsy shows no evidence of marrow involvement. Laboratory studies reveal normal renal and hepatic function. ECOG performance status is 0. Treatment options discussed include: 1) R-CVP chemotherapy (Rituximab, Cyclophosphamide, Vincristine, Prednisone) for 6 cycles, 2) Bendamustine plus Rituximab, 3) Watchful waiting. Patient elects to proceed with R-CVP regimen. Informed consent obtained. Treatment scheduled to begin next week.", "bates": "Adkisson_Patricia_2022.02.16_TxOncoHx_UCM_111_p1"}
                ],
                "record_index": [
                    {"facility": "University Iowa Hospital", "bates_range": "Adkisson_Patricia_UIH_178-185", "date_range": "10/31/2019", "description": "Comprehensive medical history documentation including past medical history, surgical history, family history, social history, and review of systems. Contains detailed physical examination findings, vital signs, laboratory results, and treatment plans. Includes documentation of chronic medical conditions and medication reconciliation."},
                    {"facility": "University Medical Center", "bates_range": "Adkisson_Patricia_UCM_111-366", "date_range": "04/01/2021 - 02/16/2022", "description": "Complete pathology reports including tissue biopsy results, immunohistochemical studies, flow cytometry analysis, and molecular studies. Contains detailed microscopic descriptions, differential diagnoses, and final pathological diagnoses. Includes oncology consultation notes, treatment planning documentation, and follow-up visit records with comprehensive physical examinations and laboratory results."},
                    {"facility": "Medical Center", "bates_range": "Adkisson_Patricia_MedCenter_146-111", "date_range": "01/17/2022 - 02/16/2022", "description": "Oncology follow-up visits and treatment planning documentation. Contains detailed physical examinations, vital signs, laboratory studies including CBC and comprehensive metabolic panel, imaging results interpretation, and treatment recommendations. Includes informed consent documentation and treatment scheduling information."}
                ],
                "qualifying_diagnosis_table": [
                    {
                        "diagnosis": "Bladder carcinoma - Morphologically low grade, mature B-Cell lymphoma, stage IV marginal zone",
                        "dx_reference": "Adkisson_Patricia_2021.04.01_PathReport_UCM_321_p1",
                        "treatment": "02/10/2022-06/02/2022 R-CVP x 6",
                        "tx_reference": "Adkisson_Patricia_2022.02.16_TxOncoHx_UCM_111_p1"
                    }
                ]
            }

        # Configure Gemini
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel("gemini-2.0-flash-exp")  # Latest Gemini 2.0 Flash with higher limits
        
        # Prepare all content in one go (no batching needed with 1M token context)
        files_content = ""
        for i, file_data in enumerate(files_data, 1):
            file_name = file_data['name']
            page_number = file_data.get('page_number', 1)
            batch_info = extract_batch_info(file_name)
            files_content += f"\n\n--- FILE {i}: {file_name} ({batch_info}) PAGE {page_number} ---\n"
            files_content += f"File Path: {file_data.get('path', 'N/A')}\n"
            files_content += f"Content:\n{file_data['content']}"
        
        logger.info(f"Processing {len(files_data)} files with Gemini 2.0 Flash (1M token context)")
        
        # Single API call with all content and retry logic
        max_retries = 3
        retry_delay = 60  # Start with 60 seconds
        
        for attempt in range(max_retries):
            try:
                response = model.generate_content(
                    [
                        {"role": "user", "parts": [MASTER_PROMPT]},
                        {"role": "user", "parts": [f"Please analyze these medical files and create a chronology summary. Use the actual file names provided for Bates number references:\n{files_content}"]}
                    ],
                    generation_config=genai.types.GenerationConfig(
                        temperature=0.1,
                        max_output_tokens=8192
                    )
                )
                break  # Success, exit retry loop
                
            except Exception as e:
                error_str = str(e)
                if "429" in error_str and "quota" in error_str.lower():
                    if attempt < max_retries - 1:
                        logger.warning(f"Rate limit hit, waiting {retry_delay} seconds before retry {attempt + 1}/{max_retries}")
                        time.sleep(retry_delay)
                        retry_delay *= 2  # Exponential backoff
                    else:
                        logger.error("Max retries reached for rate limiting")
                        raise
                else:
                    # Non-rate-limit error, don't retry
                    raise
        
        # Parse the response
        response_content = response.text
        
        # Try to extract JSON from the response
        try:
            # Look for JSON block in the response
            if "```json" in response_content:
                json_start = response_content.find("```json") + 7
                json_end = response_content.find("```", json_start)
                json_content = response_content[json_start:json_end].strip()
            else:
                json_content = response_content
            
            chronology_data = json.loads(json_content)
            return chronology_data
        
        except json.JSONDecodeError:
            # If JSON parsing fails, return structured response
            return {
                "meta_table": {
                    "patient_name": "Not stated",
                    "date_of_birth": "Not stated",
                    "date_of_diagnosis": "Not stated",
                    "age_at_diagnosis": "Not stated",
                    "primary_diagnosis": "Not stated",
                    "attorney": "Not stated",
                    "preparation_date": datetime.now().strftime("%m/%d/%Y")
                },
                "medical_history_summary": {
                    "date_of_birth": "Not stated",
                    "date_of_diagnosis": "Not stated",
                    "age_at_diagnosis": "Not stated",
                    "t1_injury_and_secondary": "Not stated",
                    "chemotherapy_treatments": [],
                    "chronological_narrative": "Not stated",
                    "sufferings": [],
                    "other_medical_history": {
                        "prior_to_t1": [],
                        "after_t1": []
                    },
                    "risk_factors": [],
                    "additional_information": {
                        "family_history": "Not stated",
                        "social_history": "Not stated",
                        "work_history": "Not stated"
                    },
                    "missing_records": "Not stated"
                },
                "chronological_records": [],
                "record_index": [],
                "qualifying_diagnosis_table": [],
                "raw_response": response_content
            }
    
    except Exception as e:
        logger.error(f"Error processing with Gemini: {str(e)}")
        raise

# Comment out the old OpenAI function for reference
"""
def process_files_with_openai_parallel_stitch(files_data: list, use_mock: bool = False) -> dict:
    model = "gpt-4o"
    max_tokens_per_batch = 15000  # Reduced to stay well under 30k TPM limit

    batches = batch_files(files_data, max_tokens_per_batch, model)

    # 1. Process batches sequentially with delays to avoid rate limiting
    batch_outputs = []
    for idx, batch in enumerate(batches):
        logger.info(f"Processing batch {idx + 1}/{len(batches)}")
        output = process_batch(batch, idx)
        batch_outputs.append(output)
        
        # Add delay between batches to allow TPM to reset (1 minute = 60 seconds)
        if idx < len(batches) - 1:  # Don't delay after the last batch
            logger.info(f"Waiting 60 seconds before next batch to reset TPM...")
            time.sleep(60)

    # 2. Stitch together - limit prompt size to avoid TPM issues
    # Prepare a prompt for the final call
    summaries = []
    for idx, output in enumerate(batch_outputs, 1):
        summaries.append(f"--- PARTIAL SUMMARY {idx} ---\n{output}")

    stitch_prompt = (
        "You are an expert medical records analyst. Here are several partial medical chronology summaries, each in the required JSON format. "
        "Please merge them into a single, unified chronology in the same JSON format, ensuring all information is included, deduplicated, and in correct chronological order.\n\n"
        + "\n\n".join(summaries)
    )

    # Check if stitching prompt is too large and truncate if needed
    stitch_tokens = num_tokens_from_string(stitch_prompt, model)
    max_stitch_tokens = 20000  # Leave room for 4k output tokens
    
    if stitch_tokens > max_stitch_tokens:
        logger.warning(f"Stitching prompt too large ({stitch_tokens} tokens), truncating to {max_stitch_tokens}")
        # Truncate by taking first part of each summary
        truncated_summaries = []
        for summary in summaries:
            # Take roughly equal portion of each summary
            target_length = max_stitch_tokens // len(summaries)
            truncated_summary = summary[:target_length] + "..." if len(summary) > target_length else summary
            truncated_summaries.append(truncated_summary)
        
        stitch_prompt = (
            "You are an expert medical records analyst. Here are several partial medical chronology summaries (truncated due to size limits), each in the required JSON format. "
            "Please merge them into a single, unified chronology in the same JSON format, ensuring all information is included, deduplicated, and in correct chronological order.\n\n"
            + "\n\n".join(truncated_summaries)
        )

    # Add delay before final stitching call to ensure TPM is reset
    logger.info("Waiting 60 seconds before final stitching call to reset TPM...")
    time.sleep(60)

    import openai
    openai.api_key = OPENAI_API_KEY
    final_response = openai.ChatCompletion.create(
        model=model,
        messages=[
            {"role": "system", "content": MASTER_PROMPT},
            {"role": "user", "content": stitch_prompt}
        ],
        max_tokens=4000,
        temperature=0.1
    )
    final_content = final_response['choices'][0]['message']['content']

    # Try to extract JSON
    try:
        if "```json" in final_content:
            json_start = final_content.find("```json") + 7
            json_end = final_content.find("```", json_start)
            json_content = final_content[json_start:json_end].strip()
        else:
            json_content = final_content

        chronology_data = json.loads(json_content)
        return chronology_data
    except json.JSONDecodeError:
        return {"raw_response": final_content}
"""

def create_chronology_document(chronology_data: Dict, patient_name: str = "Unknown") -> bytes:
    """Create Word document using the template and chronology data"""
    try:
        # Load the template
        template_path = os.path.join(os.path.dirname(__file__), "Medical_Chronology_Template.docx")
        doc = Document(template_path)
        
        # Fill in the document based on the chronology data
        meta_table = chronology_data.get("meta_table", {})
        
        # Table 0: Meta Table
        if len(doc.tables) > 0:
            meta_table_obj = doc.tables[0]
            # Fill in meta information
            if len(meta_table_obj.rows) >= 5:
                # TO: field
                if len(meta_table_obj.rows[0].cells) >= 2:
                    meta_table_obj.rows[0].cells[1].text = str(meta_table.get("attorney", "Litigation Team"))
                # FROM: field
                if len(meta_table_obj.rows[1].cells) >= 2:
                    meta_table_obj.rows[1].cells[1].text = "Medical Records Analyst"
                # DATE: field
                if len(meta_table_obj.rows[2].cells) >= 2:
                    meta_table_obj.rows[2].cells[1].text = str(meta_table.get("preparation_date", datetime.now().strftime("%m/%d/%Y")))
                # TRACK 1 DISEASE: field
                if len(meta_table_obj.rows[3].cells) >= 2:
                    meta_table_obj.rows[3].cells[1].text = str(meta_table.get("primary_diagnosis", "Medical Chronology"))
                # PLAINTIFF: field
                if len(meta_table_obj.rows) >= 5 and len(meta_table_obj.rows[4].cells) >= 2:
                    meta_table_obj.rows[4].cells[1].text = str(meta_table.get("patient_name", patient_name))
        
        # Helper function to safely join lists that might contain non-string items
        def safe_join(items, separator=", "):
            if not items:
                return "None"
            if isinstance(items, str):
                return items
            if isinstance(items, list):
                # Convert all items to strings and filter out empty ones
                str_items = [str(item) for item in items if item]
                return separator.join(str_items) if str_items else "None"
            return str(items)
        
        # Table 1: Summary of Medical History
        if len(doc.tables) > 1:
            details_table = doc.tables[1]
            medical_history = chronology_data.get("medical_history_summary", {})
            
            # Fill in details starting from row 1 (skip header)
            details_rows = [
                ("Date of Birth", str(medical_history.get("date_of_birth", "Not stated"))),
                ("Date of Diagnosis", str(medical_history.get("date_of_diagnosis", "Not stated"))),
                ("Age at Diagnosis", str(medical_history.get("age_at_diagnosis", "Not stated"))),
                ("T1 Injury + Other Related Secondary Diagnoses", str(medical_history.get("t1_injury_and_secondary", "Not stated"))),
                ("Chemotherapy Treatments", safe_join(medical_history.get("chemotherapy_treatments", []))),
                ("Chronological Narrative", str(medical_history.get("chronological_narrative", "Not stated"))),
                ("Sufferings", safe_join(medical_history.get("sufferings", []), "; ")),
                ("Other Medical History - Prior to T1", safe_join(medical_history.get("other_medical_history", {}).get("prior_to_t1", []), "; ")),
                ("Other Medical History - After T1", safe_join(medical_history.get("other_medical_history", {}).get("after_t1", []), "; ")),
                ("Risk Factors", safe_join(medical_history.get("risk_factors", []), "; ")),
                ("Family History", str(medical_history.get("additional_information", {}).get("family_history", "Not stated"))),
                ("Social History", str(medical_history.get("additional_information", {}).get("social_history", "Not stated"))),
                ("Work History", str(medical_history.get("additional_information", {}).get("work_history", "Not stated"))),
                ("Missing Records", str(medical_history.get("missing_records", "None")))
            ]
            
            for i, (description, details) in enumerate(details_rows):
                if i + 1 < len(details_table.rows):
                    row = details_table.rows[i + 1]
                    if len(row.cells) >= 2:
                        row.cells[0].text = str(description)
                        row.cells[1].text = str(details)
        
        # Table 2: Qualifying Diagnosis Table
        if len(doc.tables) > 2:
            qd_table = doc.tables[2]
            qualifying_diagnosis = chronology_data.get("qualifying_diagnosis_table", [])

            # Clear existing rows except header
            while len(qd_table.rows) > 1:
                qd_table._element.remove(qd_table.rows[1]._element)

            for entry in qualifying_diagnosis:
                new_row = qd_table.add_row()
                new_row.cells[0].text = str(entry.get("diagnosis", ""))
                
                # Format dx_reference
                dx_ref = entry.get("dx_reference", "")
                if dx_ref and "_p" in dx_ref:
                    parts = dx_ref.split("_p")
                    if len(parts) == 2:
                        file_part = parts[0]
                        page_num = parts[1]
                        formatted_dx_ref = format_document_reference(file_part, int(page_num))
                        new_row.cells[1].text = formatted_dx_ref
                    else:
                        new_row.cells[1].text = dx_ref
                else:
                    new_row.cells[1].text = dx_ref
                
                new_row.cells[2].text = str(entry.get("treatment", ""))
                
                # Format tx_reference
                tx_ref = entry.get("tx_reference", "")
                if tx_ref and "_p" in tx_ref:
                    parts = tx_ref.split("_p")
                    if len(parts) == 2:
                        file_part = parts[0]
                        page_num = parts[1]
                        formatted_tx_ref = format_document_reference(file_part, int(page_num))
                        new_row.cells[3].text = formatted_tx_ref
                    else:
                        new_row.cells[3].text = tx_ref
                else:
                    new_row.cells[3].text = tx_ref
        else:
            logger.warning("Qualifying Diagnosis table (Table 2) not found in document")
        
        # Table 3: Chronological Records
        if len(doc.tables) > 3:
            records_table = doc.tables[3]
            chronological_records = chronology_data.get("chronological_records", [])
            
            # Clear existing rows except header
            while len(records_table.rows) > 1:
                records_table._element.remove(records_table.rows[1]._element)
            
            # Add rows for each chronological record
            for record in chronological_records:
                if len(records_table.rows) > 0:
                    # Add new row
                    new_row = records_table.add_row()
                    if len(new_row.cells) >= 4:
                        new_row.cells[0].text = str(record.get("date", ""))
                        new_row.cells[1].text = str(record.get("provider_facility", ""))
                        new_row.cells[2].text = str(record.get("summary", ""))
                        
                        # Format the Bates reference for the document
                        bates_reference = record.get("bates", "")
                        if bates_reference:
                            # Try to extract file name and page number from the Bates reference
                            # Format: "FileName_BatchInfo_pX" -> "FileName-Page X"
                            if "_p" in bates_reference:
                                parts = bates_reference.split("_p")
                                if len(parts) == 2:
                                    file_part = parts[0]
                                    page_num = parts[1]
                                    # Format as DOCUMENT_NAME-Page X
                                    formatted_ref = format_document_reference(file_part, int(page_num))
                                    new_row.cells[3].text = formatted_ref
                                else:
                                    new_row.cells[3].text = bates_reference
                            else:
                                new_row.cells[3].text = bates_reference
                        else:
                            new_row.cells[3].text = ""
        
        # Table 4: Record Index
        if len(doc.tables) > 4:
            index_table = doc.tables[4]
            record_index = chronology_data.get("record_index", [])
            
            # Clear existing rows except header
            while len(index_table.rows) > 1:
                index_table._element.remove(index_table.rows[1]._element)
            
            # Add rows for each record index entry
            for index_record in record_index:
                if len(index_table.rows) > 0:
                    # Add new row
                    new_row = index_table.add_row()
                    # Check how many columns the table actually has
                    num_cols = len(new_row.cells)
                    
                    if num_cols >= 3:
                        new_row.cells[0].text = str(index_record.get("facility", ""))
                        
                        # Format bates_range for the document
                        bates_range = index_record.get("bates_range", "")
                        if bates_range and "_p" in bates_range:
                            # For ranges, extract the base file name before the page number
                            base_part = bates_range.split("_p")[0]
                            # Format as DOCUMENT_NAME (without page number for ranges)
                            formatted_range = format_document_reference(base_part, 0)
                            new_row.cells[1].text = formatted_range
                        else:
                            new_row.cells[1].text = bates_range
                        
                        new_row.cells[2].text = str(index_record.get("date_range", ""))
                        
                        # If there's a 4th column, add description
                        if num_cols >= 4:
                            new_row.cells[3].text = str(index_record.get("description", ""))
        else:
            logger.warning("Record Index table (Table 4) not found in document")
        
        # Save to bytes
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
    
    except Exception as e:
        logger.error(f"Error creating chronology document: {str(e)}")
        logger.error(f"Chronology data structure: {chronology_data}")
        raise

def upload_to_s3(file_content: bytes, s3_key: str) -> str:
    """Upload file to S3 bucket"""
    try:
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=file_content,
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Generate presigned URL for download
        download_url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET, 'Key': s3_key},
            ExpiresIn=3600  # 1 hour
        )
        
        return download_url
    
    except Exception as e:
        logger.error(f"Error uploading to S3: {str(e)}")
        raise

def lambda_handler(event, context):
    """
    Lambda handler function
    """
    try:
        # Parse the input - handle both direct invocation and API Gateway formats
        if 'body' in event:
            # API Gateway format
            if isinstance(event['body'], str):
                body = json.loads(event['body'])
            else:
                body = event['body']
        else:
            # Direct invocation format
            body = event
        
        # Get the folder path from the request
        folder_path = body.get('path', '').strip()
        max_depth = body.get('max_depth', 5)
        patient_name = body.get('patient_name', 'Unknown')
        
        logger.info(f"Processing chronology for path: '{folder_path}', max_depth: {max_depth}")
        
        # Get access token
        token = get_access_token()
        
        # Get all files from SharePoint (excluding "_CHR Claim" files)
        all_files = walk_sharepoint_path(folder_path, token, max_depth)
        
        if not all_files:
            return {
                'statusCode': 404,
                'headers': {'Content-Type': 'application/json'},
                'body': json.dumps({
                    'success': False,
                    'error': 'No files found in the specified path'
                })
            }
        
        logger.info(f"Found {len(all_files)} files to process")
        
        # Download and extract content from files
        files_data = []
        for file_info in all_files[:10]:  # Limit to first 10 files to avoid token limits
            logger.info(f"Processing file: {file_info['name']}")
            page_contents = download_file_content(file_info['download_url'], file_info['name'])
            for page in page_contents:
                files_data.append({
                    'name': file_info['name'],
                    'path': file_info['path'],
                    'content': page['text'],
                    'page_number': page['page_number']
                })
        
        # Process files with Gemini (replacing OpenAI)
        logger.info("Creating chronology with Gemini...")  
        use_mock = body.get('use_mock_ai', False)
        chronology_data = process_files_with_gemini(files_data, use_mock)
        
        # Create Word document
        logger.info("Creating Word document...")
        doc_content = create_chronology_document(chronology_data, patient_name)
        # with open("chronology.docx", "wb") as f:
        #     f.write(doc_content)
        
        
        # Upload to S3
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        s3_key = f"{folder_path}/Chronology_{patient_name}_{timestamp}.docx"
        logger.info(f"Uploading document to S3: bucket={S3_BUCKET}, key={s3_key}")
        download_url = upload_to_s3(doc_content, s3_key)
        logger.info(f"Document uploaded successfully. Download URL: {download_url}")
        
        # Call back to update the ingestion status if callback URL is provided
        callback_url = body.get('callback_url')
        if callback_url:
            try:
                callback_payload = {
                    "status": "COMPLETED",
                    "download_url": f"s3://{S3_BUCKET}/{s3_key}",
                    "processed_files_count": len(files_data)
                }
                
                logger.info(f"Calling back to: {callback_url}")
                callback_response = requests.patch(
                    callback_url,
                    json=callback_payload,
                    headers={'Content-Type': 'application/json'},
                    timeout=30
                )
                
                if callback_response.status_code == 200:
                    logger.info("Successfully updated ingestion status via callback")
                else:
                    logger.warning(f"Callback failed with status {callback_response.status_code}: {callback_response.text}")
                    
            except Exception as callback_error:
                logger.error(f"Error calling callback URL: {str(callback_error)}")
                # Don't fail the main process if callback fails
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({
                'success': True,
                'message': f'Chronology created successfully for {len(files_data)} files',
                'download_url': download_url,
                's3_bucket': S3_BUCKET,
                's3_key': s3_key,
                'files_processed': len(files_data),
                'chronology_data': chronology_data
            }, indent=2)
        }
        
    except Exception as e:
        logger.error(f"Lambda handler error: {str(e)}")
        
        # Call back to update the ingestion status to FAILED if callback URL is provided
        try:
            # Parse the input to get callback URL
            if 'body' in event:
                if isinstance(event['body'], str):
                    body = json.loads(event['body'])
                else:
                    body = event['body']
            else:
                body = event
            
            callback_url = body.get('callback_url')
            if callback_url:
                callback_payload = {
                    "status": "FAILED",
                    "error_message": str(e)
                }
                
                logger.info(f"Calling back to update status to FAILED: {callback_url}")
                callback_response = requests.patch(
                    callback_url,
                    json=callback_payload,
                    headers={'Content-Type': 'application/json'},
                    timeout=30
                )
                
                if callback_response.status_code == 200:
                    logger.info("Successfully updated ingestion status to FAILED via callback")
                else:
                    logger.warning(f"Callback failed with status {callback_response.status_code}: {callback_response.text}")
                    
        except Exception as callback_error:
            logger.error(f"Error calling callback URL for failure: {str(callback_error)}")
        
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({
                'success': False,
                'error': str(e)
            })
        } 